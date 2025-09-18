"""
DOCX Document Processor for First Cars booking extraction
Uses python-docx to extract paragraphs and tables from Word documents
"""

import os
import logging
from typing import Dict, List, Optional, Any, Tuple
from docx import Document

# Import our AI agents
from unified_email_processor import UnifiedEmailProcessor
from structured_email_agent import StructuredExtractionResult
from car_rental_ai_agent import BookingExtraction

logger = logging.getLogger(__name__)

class DocxDocumentProcessor:
    """Document processor specifically for DOCX files using python-docx"""
    
    def __init__(self, openai_api_key: str = None):
        """
        Initialize DOCX document processor
        """
        self.openai_api_key = openai_api_key or os.getenv('OPENAI_API_KEY')
        self.email_processor = UnifiedEmailProcessor(self.openai_api_key)
        
        logger.info("DOCX Document Processor initialized")
    
    def parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse DOCX file content and extract paragraphs and tables
        
        Args:
            file_content: Raw DOCX file bytes
            filename: Original filename
            
        Returns:
            Dictionary containing paragraphs and tables
        """
        try:
            # Save bytes to temporary file for python-docx
            temp_file_path = f"temp_{filename}"
            with open(temp_file_path, 'wb') as temp_file:
                temp_file.write(file_content)
            
            # Parse document
            doc = Document(temp_file_path)
            
            # Extract paragraphs
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Clean up temporary file
            try:
                os.remove(temp_file_path)
            except:
                pass
            
            result = {
                "paragraphs": paragraphs,
                "tables": tables
            }
            
            logger.info(f"DOCX parsing successful: {len(paragraphs)} paragraphs, {len(tables)} tables")
            return result
            
        except Exception as e:
            logger.error(f"DOCX parsing failed for {filename}: {str(e)}")
            # Clean up temporary file on error
            try:
                os.remove(temp_file_path)
            except:
                pass
            raise
    
    def process_document(self, file_content: bytes, filename: str) -> StructuredExtractionResult:
        """
        Process a DOCX document using the AI agent
        
        Args:
            file_content: Raw DOCX file bytes
            filename: Original filename
            
        Returns:
            StructuredExtractionResult with extracted bookings
        """
        logger.info(f"Processing DOCX document: {filename}")
        
        try:
            # Step 1: Extract content using python-docx
            parsed_data = self.parse_docx(file_content, filename)
            
            # Step 2: Convert extracted content to text for AI processing
            document_text = self._format_extracted_content(parsed_data, filename)
            
            if not document_text.strip():
                return self._create_error_result("No content extracted from document", filename)
            
            logger.info(f"Extracted {len(document_text)} characters from {filename}")
            
            # Step 3: Process with AI using enhanced context
            result = self._process_with_ai(document_text, filename, parsed_data)
            
            # Step 4: Add metadata
            result.extraction_method = f"docx_ai_processing"
            result.processing_notes = f"DOCX+AI processing: {filename}. Paragraphs: {len(parsed_data['paragraphs'])}, Tables: {len(parsed_data['tables'])}"
            
            # Step 5: Add extracted content to additional_info for each booking
            for booking in result.bookings:
                document_info = f"Document: {filename}\\nParagraphs: {len(parsed_data['paragraphs'])}\\nTables: {len(parsed_data['tables'])}\\nExtracted: {document_text[:800]}{'...' if len(document_text) > 800 else ''}"
                
                if booking.additional_info:
                    booking.additional_info = f"{booking.additional_info}\\n\\n{document_info}"
                else:
                    booking.additional_info = document_info
            
            logger.info(f"DOCX processing completed: {filename}, found {len(result.bookings)} booking(s)")
            return result
            
        except Exception as e:
            logger.error(f"DOCX processing failed for {filename}: {str(e)}")
            return self._create_error_result(str(e), filename)
    
    def _format_extracted_content(self, parsed_data: Dict[str, Any], filename: str) -> str:
        """
        Format extracted paragraphs and tables into readable text for AI processing
        """
        content_parts = []
        
        # Add document header
        content_parts.append(f"DOCUMENT: {filename}")
        content_parts.append("="*50)
        
        # Add paragraphs
        if parsed_data['paragraphs']:
            content_parts.append("\\nPARAGRAPHS:")
            content_parts.append("-" * 20)
            for i, paragraph in enumerate(parsed_data['paragraphs'], 1):
                content_parts.append(f"{i}. {paragraph}")
        
        # Add tables
        if parsed_data['tables']:
            content_parts.append("\\nTABLES:")
            content_parts.append("-" * 20)
            for table_idx, table in enumerate(parsed_data['tables'], 1):
                content_parts.append(f"\\nTable {table_idx}:")
                for row_idx, row in enumerate(table):
                    row_text = " | ".join(row)
                    content_parts.append(f"Row {row_idx + 1}: {row_text}")
        
        return "\\n".join(content_parts)
    
    def _process_with_ai(self, document_text: str, filename: str, parsed_data: Dict) -> StructuredExtractionResult:
        """
        Process extracted text with AI using enhanced context
        """
        # Create enhanced context for DOCX documents
        document_context = f"""
DOCUMENT ANALYSIS CONTEXT:
Source: {filename} (DOCX Document)
Processing Method: Python-DOCX + AI Processing
Document Type: Word Document with Booking/Reservation Information
Content Structure: {len(parsed_data['paragraphs'])} paragraphs, {len(parsed_data['tables'])} tables

EXTRACTED CONTENT FROM DOCX:
{document_text}

COMPREHENSIVE PROCESSING INSTRUCTIONS:

MULTIPLE BOOKING DETECTION:
- Analyze for multiple SEPARATE bookings (different dates/passengers/routes)
- Each unique DATE requires separate booking (17th & 18th Sept = 2 bookings)
- Table rows typically represent separate bookings
- Different passengers on different dates = separate bookings
- Round trips with overnight stays = separate outbound & return bookings
- Multi-day requirements = separate booking per day
- Document sections may contain multiple booking requirements

ZERO DATA LOSS POLICY:
- Extract EVERY piece of information from the Word document
- Driver names, contact numbers, special instructions, VIP requirements
- Corporate details, billing information, payment methods, authorization codes
- Vehicle preferences, cleanliness requirements, timing flexibility
- Emergency contacts, alternate arrangements, backup information
- Table headers, footnotes, document annotations, approval signatures
- If data doesn't fit standard fields, put in 'remarks' or 'additional_info'

CITY STANDARDIZATION:
- Extract only CITY names for from_location and to_location
- Map suburbs to cities (Jogeshwari → Mumbai, Andheri → Mumbai)
- Full addresses go in reporting_address and drop_address

TIME EXTRACTION:
- Extract exact times from document (7:43, 7:10, 7:53)
- Do not round times during extraction

Table data often contains structured booking information with multiple records.
Apply all business rules for vehicle standardization, time normalization, and route handling.
"""
        
        # Process with our unified email processor
        result = self.email_processor.process_email(document_context)
        return result
    
    def _create_error_result(self, error_message: str, filename: str = "") -> StructuredExtractionResult:
        """Create error result for failed processing"""
        error_booking = BookingExtraction(
            remarks=f"DOCX processing failed: {error_message}",
            additional_info=f"Failed to process: {filename}",
            confidence_score=0.0
        )
        
        return StructuredExtractionResult(
            bookings=[error_booking],
            total_bookings_found=0,
            extraction_method="docx_processing_error",
            confidence_score=0.0,
            processing_notes=f"DOCX processing error for {filename}: {error_message}"
        )
    
    def validate_file(self, filename: str, file_size: int, max_size: int = 10 * 1024 * 1024) -> Tuple[bool, str]:
        """Validate uploaded DOCX file"""
        if not filename:
            return False, "No filename provided"
        
        # Check file extension
        if not filename.lower().endswith('.docx'):
            return False, f"Invalid file type. Expected .docx, got: {filename.split('.')[-1]}"
        
        # Check file size
        if file_size > max_size:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB (max: {max_size / (1024*1024):.1f}MB)"
        
        return True, "File is valid"
    
    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file types"""
        return ['docx']

# Test function
def test_docx_processor():
    """Test the DOCX processor with a sample file"""
    try:
        processor = DocxDocumentProcessor()
        
        # Test with the sample file mentioned in your message
        test_file = "First Cars Booking -11-Sep-25  Deepu Sir.docx"
        
        if os.path.exists(test_file):
            with open(test_file, 'rb') as f:
                file_content = f.read()
            
            # Parse the document
            parsed_data = processor.parse_docx(file_content, test_file)
            print("Paragraphs:", parsed_data["paragraphs"])
            print("\\nTables:", parsed_data["tables"])
            
            # Process with AI
            result = processor.process_document(file_content, test_file)
            print("\\nExtracted Bookings:")
            for i, booking in enumerate(result.bookings, 1):
                print(f"Booking {i}: {booking.passenger_name}, {booking.vehicle_group}, {booking.start_date}")
        else:
            print(f"Test file {test_file} not found")
            
    except Exception as e:
        print(f"Test failed: {str(e)}")
        print("Make sure to set OPENAI_API_KEY environment variable")

if __name__ == "__main__":
    test_docx_processor()
